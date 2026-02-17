"""
Script para generar documento Word del proyecto Admin Dashboard
Pollería "Crazy Chicken"
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Instalando python-docx...")
    import subprocess

    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_custom(doc, text, level=1):
    """Agrega un encabezado con formato personalizado"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph_formatted(doc, text, bold=False, italic=False):
    """Agrega un párrafo con formato"""
    p = doc.add_paragraph(text)
    if bold or italic:
        run = p.runs[0]
        run.bold = bold
        run.italic = italic
    return p


def crear_documento():
    """Genera el documento Word completo"""
    doc = Document()

    # Portada
    title = doc.add_heading("SISTEMA DE ANÁLISIS DE VENTAS Y KPIs", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Pollería "Crazy Chicken"')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True

    doc.add_paragraph()
    info = doc.add_paragraph("Dashboard de Indicadores de Desempeño")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    fecha = doc.add_paragraph("Febrero 2026")
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.runs[0].italic = True

    doc.add_page_break()

    # Índice (manual)
    add_heading_custom(doc, "ÍNDICE", 1)
    doc.add_paragraph("1. Generación de un Diseño de Comprobación (Train y Test)")
    doc.add_paragraph("2. Generación de los Modelos")
    doc.add_paragraph("3. Evaluación de los Modelos")
    doc.add_paragraph("4. Revisión de los Resultados")
    doc.add_paragraph("5. Determinación del Modelo más Recomendado")
    doc.add_paragraph("6. Implementación")
    doc.add_paragraph("7. Conclusiones")
    doc.add_paragraph("8. Recomendaciones")

    doc.add_page_break()

    # 1. Diseño de Comprobación
    add_heading_custom(
        doc, "1. GENERACIÓN DE UN DISEÑO DE COMPROBACIÓN (TRAIN Y TEST)", 1
    )

    doc.add_paragraph(
        'En el sistema de análisis de ventas de la Pollería "Crazy Chicken", el diseño de '
        "comprobación se enfoca en la validación y filtrado de datos históricos para generar "
        "indicadores clave de desempeño (KPIs). A diferencia de un modelo de Machine Learning "
        "tradicional, este sistema no entrena modelos predictivos, sino que valida y procesa "
        "datos transaccionales para la toma de decisiones."
    )

    add_heading_custom(doc, "1.1 Funciones de Extracción de Datos", 2)

    doc.add_paragraph(
        "El sistema implementa diversas funciones especializadas para extraer y procesar "
        "información de ventas desde la base de datos SQL:"
    )

    # Lista de funciones
    funciones = [
        (
            "get_ventas_mensuales_2024()",
            "Extrae las ventas totales agrupadas por mes del año 2024",
        ),
        (
            "get_ticket_promedio_global_2024()",
            "Calcula el ticket promedio de todas las ventas",
        ),
        (
            "get_ventas_por_producto()",
            "Obtiene ventas segmentadas por tipo de producto",
        ),
        (
            "get_ventas_por_vendedor()",
            "Analiza el desempeño individual de cada vendedor",
        ),
        ("get_ranking_productos()", "Genera ranking de productos más vendidos"),
        (
            "get_participacion_productos()",
            "Calcula la participación porcentual de cada producto",
        ),
    ]

    for func, desc in funciones:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(func).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "1.2 Filtros Implementados", 2)

    doc.add_paragraph(
        "El sistema permite filtrar los datos mediante los siguientes criterios:"
    )

    filtros = [
        (
            "Filtro por Mes",
            "Permite analizar ventas de meses específicos o rangos de meses",
        ),
        (
            "Filtro por Tipo de Producto",
            "Segmenta análisis por categorías: pollos, bebidas, complementos, etc.",
        ),
        (
            "Filtro por Vendedor",
            "Evalúa el desempeño individual de cada miembro del equipo de ventas",
        ),
    ]

    for titulo, desc in filtros:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(titulo).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "1.3 Indicadores Calculados", 2)

    doc.add_paragraph(
        "A partir de los datos filtrados, el sistema calcula los siguientes indicadores clave:"
    )

    indicadores = [
        "Ventas Acumuladas: Total de ventas por período",
        "Metas Mensuales: Objetivos de venta establecidos",
        "Desviaciones: Diferencia entre ventas reales y metas",
        "Ranking: Clasificación de productos y vendedores por desempeño",
        "Participación Porcentual: Contribución de cada producto/vendedor al total",
        "Ticket Promedio: Valor promedio por transacción",
        "Crecimiento Mensual: Variación porcentual mes a mes",
    ]

    for ind in indicadores:
        doc.add_paragraph(ind, style="List Bullet")

    add_heading_custom(doc, "1.4 Visualización en Dashboard", 2)

    doc.add_paragraph(
        "Los datos validados se presentan en un dashboard interactivo desarrollado con Dash "
        "(Plotly), que incluye:"
    )

    visualizaciones = [
        "Gráficos de ventas acumuladas con líneas de meta",
        "Indicadores de desviación con código de colores (verde: cumplido, rojo: no cumplido)",
        "Gráficos de barras para crecimiento mensual",
        "Tablas de ranking de productos y vendedores",
        "Gráficos de participación porcentual (pie charts)",
        "Tarjetas de KPIs principales (ventas totales, ticket promedio, etc.)",
    ]

    for viz in visualizaciones:
        doc.add_paragraph(viz, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Nota importante: ").bold = True
    p.add_run(
        "En esta etapa NO se entrenan modelos de Machine Learning. El enfoque está en la "
        "validación, procesamiento y visualización de datos históricos para generar insights "
        "accionables."
    )

    add_heading_custom(doc, "1.5 Capturas Sugeridas", 2)
    doc.add_paragraph(
        "Dashboard de ventas acumuladas y desviaciones", style="List Bullet"
    )
    doc.add_paragraph("Gráficos de crecimiento mensual", style="List Bullet")
    doc.add_paragraph("Ranking de productos y vendedores", style="List Bullet")

    doc.add_page_break()

    # 2. Generación de Modelos
    add_heading_custom(doc, "2. GENERACIÓN DE LOS MODELOS", 1)

    doc.add_paragraph(
        'En el contexto del sistema de análisis de la Pollería "Crazy Chicken", el concepto de '
        '"modelo" se adapta a los indicadores de desempeño calculados. Cada KPI funciona como '
        "un modelo de evaluación que permite medir y predecir el comportamiento del negocio."
    )

    add_heading_custom(doc, "2.1 Modelos de Indicadores Implementados", 2)

    modelos = [
        (
            "Modelo de Ventas Acumuladas vs Meta",
            "Compara las ventas reales mensuales contra objetivos establecidos. Permite identificar "
            "meses de alto y bajo rendimiento.",
        ),
        (
            "Modelo de Crecimiento Mensual",
            "Calcula la variación porcentual de ventas mes a mes. Identifica tendencias de crecimiento "
            "o decrecimiento.",
        ),
        (
            "Modelo de Ticket Promedio",
            "Analiza el valor promedio por transacción. Útil para estrategias de upselling y cross-selling.",
        ),
        (
            "Modelo de Ranking de Productos",
            "Clasifica productos por volumen de ventas. Identifica productos estrella y de baja rotación.",
        ),
        (
            "Modelo de Participación Porcentual",
            "Calcula la contribución de cada producto/vendedor al total de ventas. Permite optimizar "
            "el mix de productos.",
        ),
        (
            "Modelo de Desempeño por Vendedor",
            "Evalúa el rendimiento individual de cada vendedor. Base para incentivos y capacitación.",
        ),
    ]

    for titulo, desc in modelos:
        add_heading_custom(doc, titulo, 3)
        doc.add_paragraph(desc)
        doc.add_paragraph()

    add_heading_custom(doc, "2.2 Cálculo de Indicadores", 2)

    doc.add_paragraph(
        "Los indicadores se calculan mediante consultas SQL optimizadas y procesamiento con Pandas:"
    )

    doc.add_paragraph()
    code = doc.add_paragraph("Ejemplo de cálculo de ventas mensuales:")
    code.runs[0].italic = True

    doc.add_paragraph(
        "SELECT MONTH(fecha) as mes, SUM(total) as ventas_totales\n"
        "FROM ventas\n"
        "WHERE YEAR(fecha) = 2024\n"
        "GROUP BY MONTH(fecha)\n"
        "ORDER BY mes",
        style="List Bullet",
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Los resultados se procesan con Pandas para calcular desviaciones, porcentajes y rankings, "
        "generando DataFrames que alimentan los componentes visuales del dashboard."
    )

    add_heading_custom(doc, "2.3 Utilidad para Decisiones", 2)

    doc.add_paragraph('Estos "modelos" de indicadores sirven para:')

    utilidades = [
        "Identificar productos con mayor demanda para optimizar inventario",
        "Detectar vendedores de alto rendimiento para programas de incentivos",
        "Reconocer meses de bajo rendimiento para planificar promociones",
        "Establecer metas realistas basadas en datos históricos",
        "Tomar decisiones estratégicas basadas en evidencia cuantitativa",
    ]

    for util in utilidades:
        doc.add_paragraph(util, style="List Bullet")

    add_heading_custom(doc, "2.4 Capturas Sugeridas", 2)
    doc.add_paragraph(
        "Dashboards mostrando indicadores calculados y comparativos",
        style="List Bullet",
    )
    doc.add_paragraph("Tablas de métricas resumidas", style="List Bullet")

    doc.add_page_break()

    # 3. Evaluación de Modelos
    add_heading_custom(doc, "3. EVALUACIÓN DE LOS MODELOS", 1)

    doc.add_paragraph(
        "La evaluación de los modelos de indicadores se realiza comparando los resultados obtenidos "
        'contra las metas establecidas por la gerencia de la Pollería "Crazy Chicken". Esta evaluación '
        "permite medir la efectividad de las estrategias comerciales implementadas."
    )

    add_heading_custom(doc, "3.1 Metas Definidas", 2)

    doc.add_paragraph("El sistema trabaja con las siguientes metas de referencia:")

    metas = [
        "Meta Mensual de Ventas: Objetivo de facturación por mes",
        "Meta de Ticket Promedio: Valor objetivo por transacción",
        "Meta de Participación: Porcentaje esperado de cada producto en ventas totales",
        "Meta de Crecimiento: Porcentaje de incremento mensual esperado",
    ]

    for meta in metas:
        doc.add_paragraph(meta, style="List Bullet")

    add_heading_custom(doc, "3.2 Análisis de Desviaciones", 2)

    doc.add_paragraph(
        "Para cada indicador, el sistema calcula la desviación respecto a la meta:"
    )

    doc.add_paragraph()
    formula = doc.add_paragraph("Desviación = (Valor Real - Meta) / Meta × 100%")
    formula.runs[0].italic = True
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Las desviaciones se clasifican en:")

    clasificacion = [
        "Positiva (verde): Cuando se supera la meta",
        "Neutral (amarillo): Cuando se está dentro del ±5% de la meta",
        "Negativa (rojo): Cuando no se alcanza la meta",
    ]

    for clas in clasificacion:
        doc.add_paragraph(clas, style="List Bullet")

    add_heading_custom(doc, "3.3 Métricas de Desempeño", 2)

    doc.add_paragraph(
        "Aunque no se utilizan métricas tradicionales de ML como accuracy o R², el sistema "
        "implementa métricas equivalentes para evaluar el desempeño:"
    )

    metricas = [
        (
            "Tasa de Cumplimiento",
            "Porcentaje de meses que cumplen o superan la meta de ventas",
        ),
        (
            "Desviación Promedio",
            "Promedio de desviaciones a lo largo del período analizado",
        ),
        ("Consistencia", "Variabilidad de las ventas mes a mes (desviación estándar)"),
        ("Tendencia", "Dirección del crecimiento (positiva, negativa o estable)"),
    ]

    for metrica, desc in metricas:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(metrica).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "3.4 Análisis de Tendencias", 2)

    doc.add_paragraph(
        "El sistema identifica patrones temporales que permiten anticipar comportamientos futuros:"
    )

    tendencias = [
        "Estacionalidad: Meses de mayor y menor demanda",
        "Productos en crecimiento: Aquellos con tendencia ascendente",
        "Vendedores en mejora: Personal con curva de aprendizaje positiva",
        "Alertas tempranas: Indicadores que muestran deterioro antes de alcanzar niveles críticos",
    ]

    for tend in tendencias:
        doc.add_paragraph(tend, style="List Bullet")

    add_heading_custom(doc, "3.5 Capturas Sugeridas", 2)
    doc.add_paragraph("Gráficos de cumplimiento de metas", style="List Bullet")
    doc.add_paragraph("Tablas de desviación vs meta", style="List Bullet")

    doc.add_page_break()

    # 4. Revisión de Resultados
    add_heading_custom(doc, "4. REVISIÓN DE LOS RESULTADOS", 1)

    doc.add_paragraph(
        "La revisión de resultados consiste en interpretar los indicadores calculados para "
        'extraer insights accionables que guíen la toma de decisiones en la Pollería "Crazy Chicken".'
    )

    add_heading_custom(doc, "4.1 Interpretación de Indicadores", 2)

    add_heading_custom(doc, "Cumplimiento de Metas Mensuales", 3)
    doc.add_paragraph(
        "El análisis mensual permite identificar períodos de alto y bajo rendimiento. "
        "Preguntas clave a responder:"
    )

    preguntas_meses = [
        "¿Qué meses superaron la meta de ventas?",
        "¿Existen patrones estacionales (ej: diciembre con mayores ventas)?",
        "¿Qué meses requieren estrategias de recuperación?",
    ]

    for preg in preguntas_meses:
        doc.add_paragraph(preg, style="List Bullet")

    add_heading_custom(doc, "Productos Destacados", 3)
    doc.add_paragraph(
        "El ranking de productos revela cuáles son los más demandados y cuáles necesitan "
        "impulso comercial:"
    )

    analisis_productos = [
        "Productos estrella: Mayor volumen de ventas y participación",
        "Productos en crecimiento: Tendencia positiva mes a mes",
        "Productos de baja rotación: Candidatos para promociones o descontinuación",
    ]

    for analisis in analisis_productos:
        doc.add_paragraph(analisis, style="List Bullet")

    add_heading_custom(doc, "Desempeño de Vendedores", 3)
    doc.add_paragraph("El análisis por vendedor permite:")

    analisis_vendedores = [
        "Identificar vendedores de alto rendimiento para reconocimiento",
        "Detectar necesidades de capacitación en vendedores con bajo desempeño",
        "Establecer benchmarks basados en los mejores vendedores",
    ]

    for analisis in analisis_vendedores:
        doc.add_paragraph(analisis, style="List Bullet")

    add_heading_custom(doc, "4.2 Identificación de Patrones", 2)

    doc.add_paragraph("El sistema permite identificar patrones clave:")

    patrones = [
        (
            "Patrón de Crecimiento",
            "Tendencia general de las ventas a lo largo del año. Permite proyectar ventas futuras.",
        ),
        (
            "Patrón de Ticket Promedio",
            "Evolución del valor promedio por transacción. Indica efectividad de estrategias de venta.",
        ),
        (
            "Patrón de Participación",
            "Cambios en la composición del mix de productos vendidos. Refleja preferencias del cliente.",
        ),
        (
            "Patrón de Vendedor",
            "Curvas de aprendizaje y desempeño del personal de ventas.",
        ),
    ]

    for patron, desc in patrones:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(patron).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "4.3 Análisis de Alertas y Desviaciones Críticas", 2)

    doc.add_paragraph(
        "El sistema genera alertas automáticas cuando se detectan desviaciones significativas:"
    )

    alertas = [
        "Alerta Roja: Desviación negativa > 15% respecto a la meta",
        "Alerta Amarilla: Desviación entre -15% y -5%",
        "Alerta de Tendencia: Dos o más meses consecutivos con desviación negativa",
        "Alerta de Producto: Caída > 20% en ventas de un producto específico",
    ]

    for alerta in alertas:
        doc.add_paragraph(alerta, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Estas alertas permiten a la gerencia tomar acciones correctivas de manera oportuna, "
        "antes de que los problemas se agraven."
    )

    add_heading_custom(doc, "4.4 Capturas Sugeridas", 2)
    doc.add_paragraph("Dashboards de crecimiento mensual", style="List Bullet")
    doc.add_paragraph(
        "Ranking de vendedores y productos destacados", style="List Bullet"
    )

    doc.add_page_break()

    # 5. Determinación del Modelo Recomendado
    add_heading_custom(doc, "5. DETERMINACIÓN DEL MODELO MÁS RECOMENDADO", 1)

    doc.add_paragraph(
        "Después de evaluar todos los indicadores, se seleccionan los KPIs más relevantes "
        'para la toma de decisiones estratégicas en la Pollería "Crazy Chicken".'
    )

    add_heading_custom(doc, "5.1 Criterios de Selección", 2)

    doc.add_paragraph("Los KPIs se evalúan según:")

    criterios = [
        ("Consistencia", "Estabilidad y confiabilidad de los datos"),
        ("Claridad", "Facilidad de interpretación para la gerencia"),
        ("Utilidad", "Impacto directo en decisiones estratégicas"),
        ("Accionabilidad", "Capacidad de generar acciones concretas"),
        ("Relevancia", "Alineación con objetivos del negocio"),
    ]

    for criterio, desc in criterios:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(criterio).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "5.2 KPIs Recomendados", 2)

    doc.add_paragraph(
        "Basándose en los criterios anteriores, se recomiendan los siguientes indicadores "
        "como los más valiosos:"
    )

    add_heading_custom(doc, "1. Ventas Acumuladas vs Meta (PRINCIPAL)", 3)

    justificacion_ventas = [
        "Indicador más directo del desempeño del negocio",
        "Permite evaluar cumplimiento de objetivos mes a mes",
        "Base para decisiones de inversión y expansión",
        "Facilita la comunicación con stakeholders",
    ]

    for just in justificacion_ventas:
        doc.add_paragraph(just, style="List Bullet")

    add_heading_custom(doc, "2. Participación de Productos (ESTRATÉGICO)", 3)

    justificacion_participacion = [
        "Identifica productos que generan mayor valor",
        "Permite optimizar inventario y cadena de suministro",
        "Guía estrategias de marketing y promociones",
        "Ayuda a definir el mix óptimo de productos",
    ]

    for just in justificacion_participacion:
        doc.add_paragraph(just, style="List Bullet")

    add_heading_custom(doc, "3. Crecimiento Mensual (TENDENCIAL)", 3)

    justificacion_crecimiento = [
        "Muestra la dirección del negocio (crecimiento o contracción)",
        "Permite proyecciones a corto y mediano plazo",
        "Identifica impacto de estrategias implementadas",
        "Alerta temprana de problemas o oportunidades",
    ]

    for just in justificacion_crecimiento:
        doc.add_paragraph(just, style="List Bullet")

    add_heading_custom(doc, "4. Desempeño por Vendedor (OPERATIVO)", 3)

    justificacion_vendedor = [
        "Permite gestión efectiva del equipo de ventas",
        "Base para programas de incentivos y bonificaciones",
        "Identifica necesidades de capacitación",
        "Facilita la planificación de recursos humanos",
    ]

    for just in justificacion_vendedor:
        doc.add_paragraph(just, style="List Bullet")

    add_heading_custom(doc, "5.3 Dashboard Integrado Recomendado", 2)

    doc.add_paragraph(
        "Se recomienda un dashboard principal que integre estos cuatro KPIs, organizado en:"
    )

    dashboard_sections = [
        "Panel Superior: Ventas acumuladas vs meta (gráfico de líneas con área)",
        "Panel Izquierdo: Participación de productos (gráfico de dona/pie)",
        "Panel Central: Crecimiento mensual (gráfico de barras)",
        "Panel Derecho: Ranking de vendedores (tabla con indicadores visuales)",
    ]

    for section in dashboard_sections:
        doc.add_paragraph(section, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Este dashboard proporciona una vista completa del negocio en una sola pantalla, "
        "facilitando la toma de decisiones rápidas e informadas."
    )

    add_heading_custom(doc, "5.4 Capturas Sugeridas", 2)
    doc.add_paragraph(
        "Dashboard comparativo con los KPIs seleccionados", style="List Bullet"
    )

    doc.add_page_break()

    # 6. Implementación
    add_heading_custom(doc, "6. IMPLEMENTACIÓN", 1)

    doc.add_paragraph(
        'El sistema de análisis de ventas de la Pollería "Crazy Chicken" está implementado '
        "como una aplicación web moderna utilizando tecnologías Python."
    )

    add_heading_custom(doc, "6.1 Arquitectura del Proyecto", 2)

    doc.add_paragraph("El proyecto utiliza la siguiente arquitectura tecnológica:")

    add_heading_custom(doc, "Stack Tecnológico", 3)

    stack = [
        (
            "Dash (Plotly)",
            "Framework para crear aplicaciones web interactivas con Python",
        ),
        ("SQL Server", "Base de datos relacional para almacenamiento de transacciones"),
        ("Pandas", "Procesamiento y análisis de datos"),
        ("Plotly", "Generación de gráficos interactivos"),
        ("Python 3.x", "Lenguaje de programación principal"),
    ]

    for tech, desc in stack:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(tech).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "6.2 Estructura de Carpetas y Módulos", 2)

    doc.add_paragraph("El proyecto está organizado en la siguiente estructura:")

    doc.add_paragraph()
    estructura = doc.add_paragraph(
        "Admin Dashboard/\n"
        "├── app/\n"
        "│   ├── pages/          # Páginas del dashboard\n"
        "│   ├── components/     # Componentes reutilizables\n"
        "│   ├── services/       # Lógica de negocio y acceso a datos\n"
        "│   └── server.py       # Punto de entrada de la aplicación\n"
        "├── data/               # Archivos de datos\n"
        "└── requirements.txt    # Dependencias del proyecto"
    )
    estructura.runs[0].font.name = "Courier New"
    estructura.runs[0].font.size = Pt(9)

    add_heading_custom(doc, "Descripción de Módulos", 3)

    modulos = [
        (
            "app/pages/",
            "Contiene las diferentes páginas del dashboard (ventas, productos, vendedores, etc.)",
        ),
        (
            "app/components/",
            "Componentes reutilizables como gráficos, tablas y tarjetas de KPIs",
        ),
        (
            "app/services/data_service.py",
            "Funciones de extracción y procesamiento de datos desde la base de datos",
        ),
        ("app/server.py", "Configuración del servidor Dash y enrutamiento de páginas"),
    ]

    for modulo, desc in modulos:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(modulo).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "6.3 Flujo de Datos", 2)

    doc.add_paragraph("El flujo de datos en el sistema sigue esta secuencia:")

    flujo = [
        "Usuario accede al dashboard a través del navegador",
        "Dash renderiza la página solicitada (ej: página de ventas)",
        "La página invoca funciones del data_service.py",
        "data_service.py ejecuta consultas SQL en la base de datos",
        "Los resultados se procesan con Pandas (filtros, cálculos, agregaciones)",
        "Los DataFrames se convierten en componentes visuales (gráficos, tablas)",
        "Dash renderiza los componentes en el navegador del usuario",
        "El usuario interactúa con filtros, actualizando los datos en tiempo real",
    ]

    for i, paso in enumerate(flujo, 1):
        doc.add_paragraph(f"{i}. {paso}")

    add_heading_custom(doc, "6.4 Ejemplo de Ejecución", 2)

    doc.add_paragraph(
        "Ejemplo de cómo se ejecuta una función y se muestra en el dashboard:"
    )

    doc.add_paragraph()
    code_example = doc.add_paragraph(
        "# En data_service.py\n"
        "def get_ventas_mensuales_2024():\n"
        '    query = """\n'
        "        SELECT MONTH(fecha) as mes, SUM(total) as ventas\n"
        "        FROM ventas WHERE YEAR(fecha) = 2024\n"
        "        GROUP BY MONTH(fecha)\n"
        '    """\n'
        "    df = pd.read_sql(query, connection)\n"
        "    return df\n\n"
        "# En la página\n"
        "df_ventas = get_ventas_mensuales_2024()\n"
        'fig = px.bar(df_ventas, x="mes", y="ventas")\n'
        "return dcc.Graph(figure=fig)"
    )
    code_example.runs[0].font.name = "Courier New"
    code_example.runs[0].font.size = Pt(9)

    add_heading_custom(doc, "6.5 Características Técnicas", 2)

    caracteristicas = [
        "Actualización en tiempo real mediante callbacks de Dash",
        "Filtros interactivos que actualizan todos los gráficos simultáneamente",
        "Diseño responsive que se adapta a diferentes tamaños de pantalla",
        "Caché de datos para mejorar el rendimiento",
        "Manejo de errores y validación de datos",
        "Interfaz intuitiva y fácil de usar",
    ]

    for caract in caracteristicas:
        doc.add_paragraph(caract, style="List Bullet")

    add_heading_custom(doc, "6.6 Capturas Sugeridas", 2)
    doc.add_paragraph("Estructura de carpetas del proyecto", style="List Bullet")
    doc.add_paragraph(
        "Ejemplo de ejecución de un módulo y salida de tabla o gráfico",
        style="List Bullet",
    )

    doc.add_page_break()

    # 7. Conclusiones
    add_heading_custom(doc, "7. CONCLUSIONES", 1)

    doc.add_paragraph(
        'El sistema de análisis de ventas desarrollado para la Pollería "Crazy Chicken" '
        "representa una solución integral para la gestión basada en datos. A continuación, "
        "se presentan las conclusiones principales del proyecto:"
    )

    add_heading_custom(doc, "7.1 Hallazgos Clave", 2)

    hallazgos = [
        (
            "Ventas Acumuladas",
            "El sistema permite un seguimiento preciso de las ventas mensuales, facilitando la "
            "identificación de períodos de alto y bajo rendimiento. La comparación con metas "
            "proporciona una métrica clara de éxito.",
        ),
        (
            "Cumplimiento de Metas",
            "La visualización de desviaciones permite a la gerencia tomar acciones correctivas "
            "de manera oportuna. El código de colores (verde/amarillo/rojo) facilita la "
            "interpretación rápida del desempeño.",
        ),
        (
            "Crecimiento Mensual",
            "El análisis de tendencias revela patrones estacionales y permite proyecciones "
            "informadas. La identificación temprana de tendencias negativas es crucial para "
            "la planificación estratégica.",
        ),
        (
            "Análisis de Productos",
            "El ranking y participación de productos proporciona insights valiosos para "
            "optimizar el inventario y enfocar esfuerzos de marketing en productos de alto valor.",
        ),
        (
            "Desempeño de Vendedores",
            "La evaluación individual permite una gestión efectiva del equipo de ventas, "
            "identificando tanto a los mejores vendedores como a aquellos que requieren apoyo.",
        ),
    ]

    for titulo, desc in hallazgos:
        add_heading_custom(doc, titulo, 3)
        doc.add_paragraph(desc)
        doc.add_paragraph()

    add_heading_custom(doc, "7.2 Importancia de los KPIs", 2)

    doc.add_paragraph(
        "Los indicadores clave de desempeño implementados son fundamentales para:"
    )

    importancia = [
        "Toma de decisiones basada en evidencia cuantitativa, no en intuición",
        "Establecimiento de objetivos realistas basados en datos históricos",
        "Monitoreo continuo del desempeño del negocio",
        "Identificación temprana de problemas y oportunidades",
        "Comunicación efectiva del estado del negocio a stakeholders",
        "Alineación de estrategias con resultados medibles",
    ]

    for imp in importancia:
        doc.add_paragraph(imp, style="List Bullet")

    add_heading_custom(doc, "7.3 Limitaciones Detectadas", 2)

    doc.add_paragraph("Es importante reconocer las limitaciones del sistema actual:")

    limitaciones = [
        (
            "No implementa Machine Learning predictivo",
            "El sistema se basa en análisis descriptivo de datos históricos. No genera "
            "predicciones automáticas de ventas futuras mediante algoritmos de ML.",
        ),
        (
            "Dependencia de datos históricos",
            "Los insights están limitados a la calidad y completitud de los datos históricos "
            "registrados en la base de datos.",
        ),
        (
            "Análisis retrospectivo",
            "El enfoque principal es analizar lo que ya ocurrió, no predecir lo que ocurrirá. "
            "Las proyecciones se basan en tendencias simples.",
        ),
        (
            "Actualización manual de metas",
            "Las metas deben ser configuradas manualmente por la gerencia. No hay optimización "
            "automática de objetivos.",
        ),
        (
            "Factores externos no considerados",
            "El sistema no integra variables externas como competencia, estacionalidad de "
            "insumos, eventos locales, etc.",
        ),
    ]

    for titulo, desc in limitaciones:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(titulo).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "7.4 Valor Agregado", 2)

    doc.add_paragraph(
        "A pesar de las limitaciones, el sistema proporciona un valor significativo:"
    )

    valor = [
        "Centraliza información dispersa en un solo dashboard",
        "Automatiza cálculos que antes se hacían manualmente en Excel",
        "Proporciona visualizaciones claras e intuitivas",
        "Permite análisis ad-hoc mediante filtros interactivos",
        "Reduce el tiempo de generación de reportes de días a minutos",
        "Establece una base sólida para futuras mejoras con ML",
    ]

    for val in valor:
        doc.add_paragraph(val, style="List Bullet")

    doc.add_page_break()

    # 8. Recomendaciones
    add_heading_custom(doc, "8. RECOMENDACIONES", 1)

    doc.add_paragraph(
        "Con base en el análisis realizado y las limitaciones identificadas, se presentan "
        "las siguientes recomendaciones para mejorar y expandir el sistema de análisis de "
        'la Pollería "Crazy Chicken".'
    )

    add_heading_custom(doc, "8.1 Mejoras en Visualización", 2)

    mejoras_viz = [
        (
            "Agregar gráficos de tendencia con proyecciones",
            "Implementar líneas de tendencia que muestren la proyección de ventas para los "
            "próximos meses basándose en datos históricos.",
        ),
        (
            "Implementar mapas de calor",
            "Visualizar patrones de ventas por día de la semana y hora del día para optimizar "
            "turnos y personal.",
        ),
        (
            "Dashboard ejecutivo resumido",
            "Crear una vista de alto nivel con los 5-6 KPIs más importantes para revisión "
            "rápida de la gerencia.",
        ),
        (
            "Alertas visuales mejoradas",
            "Implementar notificaciones push o emails automáticos cuando se detecten "
            "desviaciones críticas.",
        ),
        (
            "Comparativas año a año",
            "Agregar gráficos que comparen el desempeño actual con el mismo período del año anterior.",
        ),
    ]

    for titulo, desc in mejoras_viz:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(titulo).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "8.2 Nuevas Métricas Sugeridas", 2)

    nuevas_metricas = [
        "Tasa de retención de clientes (clientes recurrentes vs nuevos)",
        "Análisis de rentabilidad por producto (margen de ganancia)",
        "Tiempo promedio de atención por transacción",
        "Tasa de conversión de visitas a ventas",
        "Análisis de combos y productos complementarios (market basket analysis)",
        "Índice de satisfacción del cliente (si se implementa encuestas)",
    ]

    for metrica in nuevas_metricas:
        doc.add_paragraph(metrica, style="List Bullet")

    add_heading_custom(doc, "8.3 Implementación de Modelos Predictivos", 2)

    doc.add_paragraph(
        "Para evolucionar el sistema hacia capacidades predictivas, se recomienda:"
    )

    ml_recomendaciones = [
        (
            "Predicción de ventas con ARIMA o Prophet",
            "Implementar modelos de series temporales para predecir ventas futuras con "
            "intervalos de confianza.",
        ),
        (
            "Segmentación de clientes con K-Means",
            "Agrupar clientes por patrones de compra para personalizar estrategias de marketing.",
        ),
        (
            "Detección de anomalías",
            "Usar algoritmos de detección de outliers para identificar transacciones inusuales "
            "o fraudes.",
        ),
        (
            "Recomendación de productos",
            "Implementar sistemas de recomendación para sugerir combos o productos complementarios.",
        ),
        (
            "Optimización de inventario",
            "Usar modelos de demanda para optimizar niveles de stock y reducir desperdicios.",
        ),
    ]

    for titulo, desc in ml_recomendaciones:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(titulo).bold = True
        p.add_run(f": {desc}")

    add_heading_custom(doc, "8.4 Alertas Automáticas", 2)

    doc.add_paragraph("Implementar un sistema de alertas automáticas para:")

    alertas_auto = [
        "Notificar cuando las ventas diarias caigan por debajo del 80% de la meta",
        "Alertar sobre productos con inventario bajo basado en velocidad de venta",
        "Avisar cuando un vendedor tenga desempeño consistentemente bajo",
        "Notificar sobre oportunidades (productos con crecimiento acelerado)",
        "Recordar revisión de metas cuando se acerque fin de mes",
    ]

    for alerta in alertas_auto:
        doc.add_paragraph(alerta, style="List Bullet")

    add_heading_custom(doc, "8.5 Recomendaciones Estratégicas Basadas en KPIs", 2)

    add_heading_custom(doc, "Basado en Productos Más Vendidos", 3)

    rec_productos = [
        "Asegurar disponibilidad constante de productos estrella",
        "Crear combos que incluyan productos de alta rotación con productos de baja rotación",
        "Enfocar campañas de marketing en productos con mayor margen de ganancia",
        "Considerar descontinuar productos con ventas consistentemente bajas",
    ]

    for rec in rec_productos:
        doc.add_paragraph(rec, style="List Bullet")

    add_heading_custom(doc, "Basado en Metas Mensuales", 3)

    rec_metas = [
        "Ajustar metas basándose en estacionalidad detectada",
        "Implementar promociones en meses históricamente bajos",
        "Establecer incentivos para vendedores en meses críticos",
        "Planificar campañas de marketing con 2-3 semanas de anticipación a meses bajos",
    ]

    for rec in rec_metas:
        doc.add_paragraph(rec, style="List Bullet")

    add_heading_custom(doc, "Basado en Desempeño de Vendedores", 3)

    rec_vendedores = [
        "Crear programa de mentoring: vendedores top capacitan a vendedores nuevos",
        "Implementar bonificaciones basadas en cumplimiento de metas individuales",
        "Identificar y replicar mejores prácticas de vendedores destacados",
        "Proporcionar capacitación específica a vendedores con áreas de oportunidad",
        "Reconocer públicamente a vendedores del mes para motivar al equipo",
    ]

    for rec in rec_vendedores:
        doc.add_paragraph(rec, style="List Bullet")

    add_heading_custom(doc, "8.6 Roadmap Sugerido", 2)

    doc.add_paragraph(
        "Plan de implementación de mejoras a corto, mediano y largo plazo:"
    )

    add_heading_custom(doc, "Corto Plazo (1-3 meses)", 3)
    roadmap_corto = [
        "Agregar nuevas métricas básicas (rentabilidad, retención)",
        "Implementar alertas automáticas por email",
        "Mejorar visualizaciones existentes con proyecciones simples",
    ]
    for item in roadmap_corto:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_custom(doc, "Mediano Plazo (3-6 meses)", 3)
    roadmap_medio = [
        "Implementar modelos predictivos básicos (ARIMA para ventas)",
        "Desarrollar módulo de análisis de rentabilidad",
        "Crear dashboard móvil para consulta desde smartphones",
    ]
    for item in roadmap_medio:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_custom(doc, "Largo Plazo (6-12 meses)", 3)
    roadmap_largo = [
        "Sistema completo de recomendaciones con ML",
        "Integración con sistema de inventario en tiempo real",
        "Módulo de optimización automática de precios",
        "Análisis predictivo de demanda para planificación de compras",
    ]
    for item in roadmap_largo:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph()

    # Cierre del documento
    cierre = doc.add_paragraph(
        'La implementación de estas recomendaciones permitirá a la Pollería "Crazy Chicken" '
        "evolucionar de un sistema de análisis descriptivo a uno predictivo y prescriptivo, "
        "maximizando el valor de los datos para la toma de decisiones estratégicas."
    )
    cierre.runs[0].italic = True

    # Guardar documento
    output_path = r"c:\Users\User\Documents\Zegel\LABO V\Admin Dashboard\Informe_Sistema_Analisis_Ventas_Crazy_Chicken.docx"
    doc.save(output_path)
    print(f"\n✅ Documento generado exitosamente en:\n{output_path}")
    return output_path


if __name__ == "__main__":
    try:
        path = crear_documento()
        print(f"\n📄 El documento Word ha sido creado correctamente.")
        print(f"📍 Ubicación: {path}")
    except Exception as e:
        print(f"\n❌ Error al generar el documento: {e}")
        import traceback

        traceback.print_exc()
